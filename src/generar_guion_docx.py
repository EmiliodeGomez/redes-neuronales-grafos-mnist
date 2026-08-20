from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guion_Exposicion_12_Minutos.docx"

NAVY = "16324F"
BLUE = "2E74B5"
TEAL = "138A8A"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E7F5F4"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "667085"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), border_color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_speaker_heading(doc, name, time_text, slides):
    p = doc.add_paragraph(style="Heading 1")
    keep_with_next(p)
    run = p.add_run(name)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    meta = p.add_run(f"  |  {time_text}  |  diapositivas {slides}")
    meta.bold = False
    meta.font.size = Pt(10.5)
    meta.font.color.rgb = RGBColor.from_string(MID_GRAY)
    return p


def add_slide(doc, number, cue, paragraphs):
    h = doc.add_paragraph(style="Heading 2")
    keep_with_next(h)
    tag = h.add_run(f"DIAPOSITIVA {number}")
    tag.bold = True
    tag.font.color.rgb = RGBColor.from_string(TEAL)
    title = h.add_run(f"  {cue}")
    title.bold = False
    title.font.color.rgb = RGBColor.from_string(NAVY)
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = True


def add_transition(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, PALE_TEAL, TEAL)
    run = p.add_run("Transición: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(text)


def add_qa(doc, question, answer):
    h = doc.add_paragraph(style="Heading 2")
    keep_with_next(h)
    h.add_run(question)
    p = doc.add_paragraph(answer)
    p.paragraph_format.keep_together = True


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"

    title_style = doc.styles["Title"]
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(NAVY)
    title_style.paragraph_format.space_after = Pt(8)

    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.size = Pt(13)
    subtitle_style.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle_style.paragraph_format.space_after = Pt(18)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("1F4D78")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "BCD5105 MODELADO MATEMÁTICO  |  GUION DE EXPOSICIÓN"
    hp.style = doc.styles["Normal"]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(TEAL)

    footer = section.footer
    fp = footer.paragraphs[0]
    set_repeat_page_number(fp)

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(8)
    run = eyebrow.add_run("GUION PALABRA POR PALABRA  |  12 MINUTOS")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph("Digitalización de archivos públicos", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(
        "Clasificación de dígitos con contexto local: prototipo OCR Softmax vs. Graph-MLP",
        style="Subtitle",
    )

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    r = meta.add_run("Lead University  ·  19 de agosto de 2026\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    meta.add_run(
        "Diego Díaz · Bairon Horna · Ignacio Marín · Jordan López · "
        "Aaron Medrano · Emilio de Gomez"
    )

    agenda = doc.add_table(rows=1, cols=4)
    agenda.alignment = WD_TABLE_ALIGNMENT.CENTER
    agenda.autofit = False
    widths = [Inches(1.0), Inches(2.05), Inches(1.35), Inches(2.10)]
    headers = ["Tiempo", "Integrante", "Diapositivas", "Enfoque"]
    for idx, (cell, width, text) in enumerate(zip(agenda.rows[0].cells, widths, headers)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(agenda.rows[0])

    rows = [
        ("0:00-2:00", "Diego Díaz", "1-2", "Problema y propuesta"),
        ("2:00-4:00", "Bairon Horna", "3-4", "Pregunta, objetivos y datos"),
        ("4:00-6:00", "Ignacio Marín", "5-6", "EDA y construcción del grafo"),
        ("6:00-8:00", "Jordan López", "7-8", "Modelos y validación"),
        ("8:00-10:00", "Aaron Medrano", "9-10", "Resultados y errores"),
        ("10:00-12:00", "Emilio de Gomez", "11-13", "Sensibilidad, límites y cierre"),
    ]
    for ridx, row in enumerate(rows):
        cells = agenda.add_row().cells
        for idx, (cell, width, text) in enumerate(zip(cells, widths, row)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ridx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(9.3)
    set_table_borders(agenda)

    doc.add_paragraph()
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(5)
    set_paragraph_shading(callout, PALE_BLUE, BLUE)
    rr = callout.add_run("REGLA CLAVE PARA TODO EL EQUIPO\n")
    rr.bold = True
    rr.font.color.rgb = RGBColor.from_string(BLUE)
    callout.add_run(
        "No digan que el modelo ya funciona en municipalidades o instituciones públicas. "
        "MNIST es el banco de prueba; la aplicación costarricense propuesta requiere un piloto local, "
        "revisión humana y conservación de la imagen original."
    )

    repo = doc.add_paragraph()
    repo.paragraph_format.space_before = Pt(10)
    rr = repo.add_run("Repositorio: ")
    rr.bold = True
    repo.add_run("https://github.com/EmiliodeGomez/redes-neuronales-grafos-mnist")

    doc.add_page_break()

    add_speaker_heading(doc, "Diego Díaz", "0:00-2:00", "1-2")
    add_slide(doc, "1", "Apertura y contexto", [
        "Buenos días. Nuestro proyecto se titula Digitalización de archivos públicos: clasificación de dígitos con contexto local. El punto de partida es un problema real: muchas municipalidades e instituciones públicas conservan protocolos, formularios y expedientes en papel. Para migrarlos a sistemas digitales no basta con escanear; también hay que transcribir fechas, folios y montos con trazabilidad.",
        "Dentro de un proceso OCR, reconocer correctamente los dígitos es una tarea pequeña, pero crítica. Un error en un monto, una fecha o un número de expediente puede cambiar el significado de un registro. Por eso proponemos evaluar si la relación espacial entre los píxeles ayuda a reconocer el trazo completo.",
        "Es importante aclarar el alcance desde el inicio: no afirmamos que el modelo ya esté listo para archivos costarricenses. Utilizamos MNIST como banco de prueba controlado para comparar dos métodos y definir qué haría falta antes de un piloto institucional."
    ])
    add_slide(doc, "2", "Por qué representar un dígito como red", [
        "La idea central es simple: un dígito puede verse como una red. Cada uno de sus 784 píxeles es un nodo, y cada nodo se conecta con sus ocho vecinos inmediatos. Así, el modelo no observa únicamente intensidades aisladas; también incorpora la continuidad de curvas, diagonales y uniones del trazo.",
        "Esta representación es relevante para documentos físicos porque el ruido, la tinta tenue o una pequeña ruptura pueden afectar un píxel, mientras que el patrón de su vecindad todavía conserva información. Con este planteamiento, le doy la palabra a Bairon para explicar la pregunta de investigación, los objetivos y los datos."
    ])
    add_transition(doc, "Bairon continúa con la pregunta de investigación, los objetivos medibles y la preparación del banco de prueba.")

    add_speaker_heading(doc, "Bairon Horna", "2:00-4:00", "3-4")
    add_slide(doc, "3", "Pregunta y objetivos medibles", [
        "Nuestra pregunta es: ¿el contexto local mejora el reconocimiento de dígitos y qué falta para validarlo en archivos públicos? A partir de ella definimos cuatro objetivos medibles.",
        "Primero, representar cada dígito como un grafo regular de 784 nodos. Segundo, comparar una regresión Softmax con un Graph-MLP bajo exactamente la misma prueba. Tercero, validar con F1 macro, exactitud, precisión, exhaustividad y pérdida logarítmica. Y cuarto, estudiar la sensibilidad al nivel de regularización, la cantidad de propagación y el tamaño de los datos.",
        "El criterio principal es F1 macro, porque resume el desempeño dando el mismo peso a cada clase. La meta experimental no es declarar resuelto el OCR institucional, sino comprobar si el prototipo con contexto supera de forma consistente al modelo base en este banco de prueba."
    ])
    add_slide(doc, "4", "Datos y separación experimental", [
        "Usamos MNIST, con 70 mil imágenes de 28 por 28 píxeles. El conjunto tiene 60 mil ejemplos de entrenamiento y 10 mil de prueba oficial. Las diez clases están razonablemente balanceadas: la relación entre la clase más y menos frecuente es de 1.24.",
        "Normalizamos las intensidades al intervalo de cero a uno y conservamos cerrada la prueba oficial hasta el final. Para la validación cruzada tomamos 12 mil observaciones estratificadas del entrenamiento, con tres pliegues compartidos por ambos modelos. Fijamos la semilla 2026 para reproducibilidad.",
        "No hay valores faltantes ni imputación, y las etiquetas no participan en la propagación del grafo. También evitamos rutas personales en el código. Pero MNIST no contiene escritura costarricense; esa limitación debe acompañar cualquier interpretación. Ignacio explicará ahora cómo exploramos las imágenes y cómo construimos el grafo."
    ])
    add_transition(doc, "Ignacio muestra la evidencia visual y la corrección matemática que permite que la propagación funcione.")

    doc.add_page_break()

    add_speaker_heading(doc, "Ignacio Marín", "4:00-6:00", "5-6")
    add_slide(doc, "5", "Exploración visual", [
        "La exploración confirma dos cosas. Primero, las clases están suficientemente balanceadas para una comparación estable. Por ejemplo, el dígito uno tiene 6,742 ejemplos y el cinco tiene 5,421. Segundo, las imágenes promedio conservan una geometría clara: aparecen segmentos verticales, curvas y diagonales.",
        "Esto justifica conectar cada píxel no solo con arriba, abajo, izquierda y derecha, sino también con las cuatro diagonales. Las visualizaciones del informe incluyen muestras y distribución de clases, promedios por clase, comparación de métricas, matrices de confusión, sensibilidad y convergencia. Cada figura se interpreta en el texto y puede reproducirse desde los notebooks."
    ])
    add_slide(doc, "6", "Propagación sobre el grafo", [
        "La propagación combina la señal propia con el promedio de los vecinos. En cada paso, H de t más uno es igual a uno menos alfa por H de t, más alfa por D inversa A H de t. Usamos alfa igual a 0.35, por lo que cada nodo conserva 65 por ciento de su intensidad y recibe 35 por ciento del vecindario. Aplicamos dos pasos para incorporar contexto sin suavizar demasiado.",
        "Aquí corregimos un error importante del prototipo anterior. La operación einsum con el índice 'nn' seleccionaba la diagonal de la matriz de adyacencia. Como esa diagonal es cero, eliminaba todos los mensajes y el modelo de grafos no podía aprender. La implementación revisada suma realmente los ocho vecinos y normaliza por el grado.",
        "El resultado es una representación que mantiene el trazo original y agrega contexto local. Jordan explicará cómo usamos esa representación para comparar los dos modelos sin favorecer a ninguno."
    ])
    add_transition(doc, "Jordan presenta los supuestos de cada modelo y el esquema común de validación.")

    add_speaker_heading(doc, "Jordan López", "6:00-8:00", "7-8")
    add_slide(doc, "7", "Modelos comparados", [
        "El modelo base es una regresión Softmax. Recibe las 784 intensidades originales y aprende una frontera lineal entre las diez clases. Tiene 7,850 parámetros y regularización L2 de diez a la menos cuatro.",
        "El segundo modelo es Graph-MLP. Primero aplica la propagación local de dos pasos y después usa una capa oculta de 64 unidades con activación ReLU antes de clasificar. Tiene aproximadamente 50,900 parámetros. Ambos se entrenan durante 12 épocas, con lotes de 256 y semilla 2026.",
        "La comparación permite saber si una arquitectura con contexto y no linealidad mejora el desempeño práctico. Sin embargo, no atribuye toda la ganancia únicamente al grafo, porque Graph-MLP también tiene una capa no lineal y más parámetros. Esa es una limitación que reconocemos explícitamente."
    ])
    add_slide(doc, "8", "Validación equivalente", [
        "Para hacer la evaluación comparable, ambos modelos reciben los mismos índices en cada uno de los tres pliegues estratificados. Cada observación valida una vez, y la prueba oficial de 10 mil imágenes se usa solamente para la evaluación final.",
        "Nuestra métrica principal es F1 macro. También reportamos exactitud, precisión macro, exhaustividad macro y log-loss. Esta última es importante porque penaliza predicciones incorrectas hechas con demasiada confianza.",
        "Con esta metodología, una mejora debe aparecer en los mismos pliegues y mantenerse en la prueba cerrada. Aaron presenta los resultados y los tipos de error que todavía persisten."
    ])
    add_transition(doc, "Aaron cuantifica la mejora y muestra por qué los casos ambiguos aún requieren revisión humana.")

    doc.add_page_break()

    add_speaker_heading(doc, "Aaron Medrano", "8:00-10:00", "9-10")
    add_slide(doc, "9", "Resultados principales", [
        "En la prueba oficial, Softmax alcanza 92.30 por ciento de exactitud, 92.19 por ciento de F1 macro y una pérdida logarítmica de 0.280. Graph-MLP alcanza 95.56 por ciento de exactitud, 95.51 por ciento de F1 macro y una pérdida de 0.152.",
        "La diferencia es de 3.31 puntos porcentuales de F1 a favor de Graph-MLP. Además, la pérdida logarítmica disminuye 45.7 por ciento, lo que indica probabilidades mejor calibradas en este banco de prueba. Graph-MLP supera a Softmax en los tres de tres pliegues de validación, con desviaciones menores a 0.4 puntos.",
        "Por lo tanto, la mejora es consistente dentro de MNIST. No obstante, estos números no estiman el rendimiento en expedientes reales: miden imágenes centradas, limpias y ya segmentadas."
    ])
    add_slide(doc, "10", "Errores que permanecen", [
        "La matriz de confusión muestra que el problema no desaparece. La confusión más persistente es entre cuatro y nueve: hay 28 casos en cada dirección. También aparecen 21 ochos clasificados como cinco y 15 tres clasificados como cinco.",
        "La clase cinco mejora de 88.8 por ciento de aciertos con Softmax a 95.1 por ciento con Graph-MLP. Aun así, un sistema para archivo público no debería aceptar automáticamente toda predicción. Los casos de baja confianza deben ir a revisión humana, y la imagen original debe conservarse para poder auditar o corregir la transcripción.",
        "Emilio cerrará con el análisis de sensibilidad, las limitaciones y la propuesta concreta para llevar este trabajo a datos costarricenses."
    ])
    add_transition(doc, "Emilio conecta la evidencia experimental con una ruta responsable de implementación local.")

    add_speaker_heading(doc, "Emilio de Gomez", "10:00-12:00", "11-13")
    add_slide(doc, "11", "Sensibilidad e interpretación", [
        "El análisis de sensibilidad evita quedarnos con una sola configuración. En Softmax, cambiar L2 mantiene el desempeño en un rango estrecho, aproximadamente entre 92.51 y 92.37 por ciento. En Graph-MLP, aumentar la propagación no siempre ayuda: con T igual a cero se obtiene 92.73 por ciento en la prueba de sensibilidad, mientras que con T igual a tres baja a 92.21 por ciento.",
        "Esto indica que demasiada propagación suaviza los trazos y elimina diferencias útiles. También vemos que el modelo es más frágil cuando se reduce la cantidad de datos. La interpretación honesta es que la no linealidad aporta una parte importante de la mejora y que el suavizado fijo debe controlarse."
    ])
    add_slide(doc, "12", "Limitaciones y salvaguardas", [
        "El 95.5 por ciento en MNIST no autoriza migrar un archivo. MNIST usa dígitos centrados y limpios, no representa caligrafía costarricense, y nuestra propagación es fija; no es una GNN aprendible de extremo a extremo. Además, la comparación mezcla el efecto del contexto con el de la no linealidad y todavía no incluye una auditoría demográfica.",
        "Antes de un uso institucional proponemos cinco controles: validar por tipo y fuente documental, definir un umbral de rechazo, enviar casos dudosos a revisión humana, conservar siempre la imagen original y monitorear la deriva y el desempeño por subgrupos o fuentes. La responsabilidad final no se delega al algoritmo."
    ])
    add_slide(doc, "13", "Conclusión y siguiente paso", [
        "Concluimos que Graph-MLP mejora el F1 y la pérdida logarítmica frente a Softmax en MNIST, pero todavía no demuestra fidelidad sobre documentos públicos. El siguiente paso no es desplegar: es construir evidencia local.",
        "Recomendamos seleccionar una serie documental real de una municipalidad o institución pública, etiquetar entre 1,000 y 5,000 recortes de dígitos, repetir la comparación y definir el umbral de revisión con personal archivístico. Si los resultados se sostienen, el prototipo puede integrarse como apoyo, nunca como sustituto del control humano.",
        "Cerramos con una idea: digitalizar no significa borrar el papel; significa preservar, transcribir y poder auditar. Muchas gracias."
    ])

    doc.add_page_break()
    doc.add_paragraph("Preguntas probables y respuestas breves", style="Title")
    p = doc.add_paragraph(
        "Estas respuestas son para las diapositivas de respaldo 14 y 15. Responder en 20-30 segundos y volver siempre al alcance real del prototipo.",
        style="Subtitle",
    )
    p.paragraph_format.space_after = Pt(12)

    add_qa(doc, "¿El modelo ya funciona en municipalidades?", "No. Los resultados corresponden a MNIST. El trabajo propone la aplicación, pero antes exige datos costarricenses etiquetados, una prueba por fuente documental y revisión humana.")
    add_qa(doc, "¿Por qué usar una representación de grafo?", "Porque los píxeles forman una estructura espacial y los trazos dependen de relaciones locales. El grafo permite mezclar la intensidad de cada píxel con la de sus ocho vecinos.")
    add_qa(doc, "¿Toda la mejora viene del grafo?", "No podemos afirmarlo. Graph-MLP también añade ReLU y más parámetros. El análisis de sensibilidad sugiere que la no linealidad aporta bastante y que demasiada propagación perjudica.")
    add_qa(doc, "¿Qué pasa si el sistema se equivoca?", "Una predicción de baja confianza se rechaza y pasa a revisión humana. La imagen original se conserva para auditoría y corrección; nunca debe sobrescribirse por una salida automática.")
    add_qa(doc, "¿Cuál es el siguiente paso concreto?", "Etiquetar entre 1,000 y 5,000 recortes de una serie documental pública, separar los datos por fuente, repetir la comparación y medir el desempeño antes de cualquier piloto operativo.")
    add_qa(doc, "¿Cuál fue el error principal del notebook anterior?", "La expresión einsum con 'nn' tomaba la diagonal nula de la matriz de adyacencia y eliminaba todos los mensajes. La versión revisada suma los ocho vecinos y normaliza por el grado.")

    reminders = doc.add_paragraph("Recordatorios finales", style="Heading 1")
    keep_with_next(reminders)
    for item in (
        "Ensayar una vez con cronómetro y respetar los relevos de dos minutos.",
        "Decir 'prototipo' y 'aplicación propuesta'; no decir 'sistema implementado en gobierno'.",
        "Pronunciar las métricas con calma: 92.19% frente a 95.51% de F1 macro.",
        "Si falta tiempo, recortar ejemplos; no eliminar limitaciones, ética ni siguiente paso local.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    props = doc.core_properties
    props.title = "Guion de exposición - Digitalización de archivos públicos"
    props.subject = "Guion palabra por palabra para exposición de 12 minutos"
    props.author = "Diego Díaz; Bairon Horna; Ignacio Marín; Jordan López; Aaron Medrano; Emilio de Gomez"
    props.keywords = "digitalización, archivos públicos, OCR, MNIST, Graph-MLP, exposición"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
